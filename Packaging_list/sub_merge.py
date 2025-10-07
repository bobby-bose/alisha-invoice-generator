# Requirements:
# pip install python-docx
# pip install pydocx
# pip install xhtml2pdf

import re
import json
from docx import Document
from pydocx import PyDocX
from xhtml2pdf import pisa
import os
import copy
import collections
import collections.abc

if not hasattr(collections, 'Hashable'):
    collections.Hashable = collections.abc.Hashable


def extract_placeholders(docx_path):
    """Extract placeholders in the format «key» from a .docx file."""
    doc = Document(docx_path)
    placeholders = set()
    pattern = r'«(.*?)»'

    for para in doc.paragraphs:
        placeholders.update(re.findall(pattern, para.text))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    placeholders.update(re.findall(pattern, para.text))

    return list(placeholders)


def replace_placeholders_and_items(docx_path, data, output_docx):
    """Replace placeholders and dynamically handle repeating table rows."""
    doc = Document(docx_path)
    pattern = r'«(.*?)»'

    # Function to replace placeholders in a paragraph
    def replace_in_para(para, local_data):
        for run in para.runs:
            matches = re.findall(pattern, run.text)
            for match in matches:
                if match in local_data:
                    run.text = run.text.replace(f'«{match}»', str(local_data[match]))

    # Step 1: Replace all placeholders (non-table, single keys)
    for para in doc.paragraphs:
        replace_in_para(para, data)

    # Step 2: Replace placeholders inside tables for normal keys
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_para(para, data)

    # Step 3: Handle repeating table rows for "items"
    if "items" in data and isinstance(data["items"], list) and len(data["items"]) > 0:
        for table in doc.tables:
            base_row = None
            base_index = None

            # Find the row that contains placeholders like «item_number», etc.
            for i, row in enumerate(table.rows):
                row_text = " ".join(cell.text for cell in row.cells)
                if any(f"«{key}»" in row_text for key in data["items"][0].keys()):
                    base_row = row
                    base_index = i
                    break

            # If such a row is found — expand it dynamically
            if base_row is not None:
                for i, item in enumerate(data["items"]):
                    new_row = copy.deepcopy(base_row)
                    table._tbl.insert(base_index + i + 1, new_row._tr)
                    new_row = table.rows[base_index + i + 1]
                    for cell in new_row.cells:
                        for para in cell.paragraphs:
                            replace_in_para(para, item)

                # Remove the original placeholder row
                del table._tbl.tr_lst[base_index]

    doc.save(output_docx)
    print(f"✅ New DOCX generated: {output_docx}")


def docx_to_pdf(input_path, output_path):
    """Convert DOCX to PDF using PyDocX + xhtml2pdf."""
    html = PyDocX.to_html(input_path)
    with open(output_path, "wb") as pdf_file:
        pisa_status = pisa.CreatePDF(html, dest=pdf_file)
        if pisa_status.err:
            print(f"❌ PDF conversion failed with errors: {pisa_status.err}")
        else:
            print(f"✅ PDF created successfully: {output_path}")


if __name__ == "__main__":
    template_path = "2.docx"  # your template file
    json_path = "data.json"   # JSON data file

    if not os.path.exists(json_path):
        print(f"❌ JSON file not found at: {json_path}")
        exit(1)
    if not os.path.exists(template_path):
        print(f"❌ Template not found at: {template_path}")
        exit(1)

    # Load JSON data
    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    # Show found placeholders
    keys = extract_placeholders(template_path)
    print("🧩 Placeholders found in document:")
    for k in keys:
        print(f" - {k}")

    # Generate output paths
    output_docx = "output_filled.docx"
    output_pdf = "output_filled.pdf"

    # Perform merge
    print("\n🔄 Replacing placeholders and expanding items...")
    replace_placeholders_and_items(template_path, data, output_docx)

    # Convert to PDF
    print("📄 Converting to PDF...")
    docx_to_pdf(output_docx, output_pdf)
