from docx import Document
import copy
import json
import os
from datetime import datetime

# ---------- Get paths relative to this script/exe ----------


# ---------- Load DOCX and JSON ----------


# ---------- Function 1: Populate second table with items ----------
def populate_table_items(table, items):
    if len(table.rows) < 3:
        print("⚠️ Table does not have at least 3 rows (header, template, total)")
        return

    template_row = table.rows[1]
    total_row = table.rows[2]

    while len(table.rows) > 2:
        table._tbl.remove(table.rows[1]._tr)

    for item in items:
        new_row = copy.deepcopy(template_row._tr)
        table._tbl.insert(-1, new_row)

    for i, item in enumerate(items):
        row = table.rows[i + 1]
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    text = run.text
                    replaced_any = False
                    for key, val in item.items():
                        placeholder = f"«{key}»"
                        if placeholder in text:
                            text = text.replace(placeholder, val)
                            replaced_any = True
                    run.text = text
                    if replaced_any:
                        print(f"✔ Replaced placeholders in row {i+1}: {item}")
    print("✅ Finished populating second table with items.\n")

# ---------- Function 2: Replace other placeholders ----------
def replace_other_placeholders(doc, data):
    placeholders_found = set()
    placeholders_not_found = set()

    for table_index, table in enumerate(doc.tables):
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        text = run.text
                        replaced_any = False
                        for key, val in data.items():
                            if key == "items":
                                continue
                            placeholder = f"«{key}»"
                            if placeholder in text:
                                text = text.replace(placeholder, val)
                                replaced_any = True
                                placeholders_found.add(placeholder)
                        run.text = text
                        if not replaced_any:
                            if "«" in text and "»" in text:
                                placeholders_not_found.add(text)
    print("✔ Placeholders replaced from top-level JSON keys:", placeholders_found)
    if placeholders_not_found:
        print("⚠️ Placeholders not found or not replaced:", placeholders_not_found)
    print("✅ Finished replacing other placeholders.\n")

# ---------- Execute Functions ----------


# ---------- Save updated document ----------

def start():
    now = datetime.now()
    day = now.day
    month_name = now.strftime("%B")  # full month name
    year = now.year
    hour = now.hour
    minute = now.minute

    # Format am/pm
    am_pm = "am" if hour < 12 else "pm"
    base_dir = os.path.dirname(os.path.abspath(__file__))
    hour_12 = hour if 1 <= hour <= 12 else (hour - 12 if hour > 12 else 12)
    output_filename = f"Packaging-List-{day}{month_name}{year}-{hour_12}.{minute}{am_pm}.docx"
    output_doc = os.path.join(base_dir, output_filename)
    
    input_doc = os.path.join(base_dir, "2.docx")
    json_data = os.path.join(base_dir, "data.json")
    doc = Document(input_doc)
    data = json.load(open(json_data, "r", encoding="utf-8"))
    populate_table_items(doc.tables[1], data["items"])
    replace_other_placeholders(doc, data)
    doc.save(output_doc)
    print(f"💾 Document updated and saved as: {output_doc}")
    

    # ---------- Generate output filename dynamically ----------
    

    

    print(f"📂 Base directory: {base_dir}")
    print(f"📄 Input DOCX: {input_doc}")
    print(f"📄 JSON file: {json_data}")
    print(f"💾 Output DOCX: {output_doc}")

