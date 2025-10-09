from pymongo import MongoClient
import sys
from PyQt6.QtWidgets import (
    QApplication, QWidget, QVBoxLayout, QHBoxLayout, QLabel, QPushButton,
    QTableWidget, QTableWidgetItem, QMessageBox, QHeaderView
)
from PyQt6.QtGui import QFont, QPixmap
from PyQt6.QtCore import Qt
import os
from ZC_IN_07.one_edit import ExportFormEdit
from bson.objectid import ObjectId
import json

# ---------------- MongoDB Config ----------------
MONGO_URL = "mongodb+srv://admin:admin@cluster.rnhig2f.mongodb.net/?retryWrites=true&w=majority&appName=cluster"
DB_NAME = "mamshi"
COLLECTION_NAME = "report_two"

# ---------------- Placeholder Mapping ----------------
PLACEHOLDER_MAPPING = {
    "Invoice No": "invoice_no",
    "Exporter's Ref": "exporter_ref",
    "IEC": "iec",
    "GST No": "tax_no",
    "Consignee": "consignee_address",
    "Pre-carraige by": "pre_carriage",
    "Port Of Discharge": "port_Discharge",
    "Date": "date",
    "Buyer's order No": "order_no",
    "Buyer's order Date": "order_date",
    "Other references": "other_reference",
    "TAX Registration Number": "tax_no",
    "Country of origin of Goods": "country",
    "Country Of Final Destination": "final_destination",
    "Country Of Destination": "destination",
    "Terms Of Delivery and Payment": "terms",
    "LUT ARN NO": "lut_arn_no",
    "Vessel / Flight No": "vessel_no",
    "Final Destination": "port_destination",
    "Port Of landing": "port_loading",
    "AD Code": "ad_code",
    "Total Packages": "total_packages",
    "Total Export Value": "export_values",
    "Total GST Value": "gst_values",
    "Total Invoice Value": "total_invoice_value",
    "Units": "units",
    "Qty": "qty",
    "Rate": "rate",
    "Amount": "amount",
    "Taxable Value": "taxable_value",
    "IGST": "igst",
    "IGST Amount": "igst_amount",
    "Amount Chargeable (in words)": "amount_inwords",
    "Shipping Mark": "shipping_mark",
    "Description of Goods": "description",
    "Delivery": "delivery",
    "Contact Details": "contact_details",
}

# ---------------- Document Functions ----------------
import re
from docx import Document
from PIL import Image

def extract_placeholders(docx_path):
    doc = Document(docx_path)
    placeholders = set()
    pattern = r'«(.*?)»'
    for para in doc.paragraphs:
        placeholders.update(re.findall(pattern, para.text))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    placeholders.update(re.findall(pattern, para.text))
    return list(placeholders)

def replace_placeholders(docx_path, data, output_docx):
    doc = Document(docx_path)
    pattern = r'«(.*?)»'

    def replace_in_para(para):
        for run in para.runs:
            matches = re.findall(pattern, run.text)
            for match in matches:
                if match in data:
                    run.text = run.text.replace(f'«{match}»', str(data[match]))

    for para in doc.paragraphs:
        replace_in_para(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_para(para)
    doc.save(output_docx)

def docx_to_image(docx_path, image_path):
    doc = Document(docx_path)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"

    img = Image.new("RGB", (1200, 1600), color="white")
    from PIL import ImageDraw, ImageFont
    draw = ImageDraw.Draw(img)
    font = ImageFont.load_default()
    draw.text((50, 50), text, fill="black", font=font)
    img.save(image_path)
    return image_path

def image_to_pdf(image_path, pdf_path):
    img = Image.open(image_path)
    img = img.convert("RGB")
    img.save(pdf_path)

# ---------------- Fetch Data ----------------
def fetch_exporter_docs():
    try:
        client = MongoClient(MONGO_URL)
        db = client[DB_NAME]
        collection = db[COLLECTION_NAME]
        docs = list(collection.find({"exporter_ref": {"$exists": True}}))
        print(f"✅ {len(docs)} documents found with key 'exporter_ref'")
        return docs
    except Exception as e:
        print("❌ MongoDB Error:", str(e))
        return []

# ---------------- UI Class ----------------
class ExportTable(QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Exporter Data Viewer")
        self.setGeometry(100, 50, 1300, 900)

        self.setStyleSheet("""
            QWidget { background-color: #ffffff; font-family: Arial; }
            QLabel { color: #1c1e21; font-weight: bold; }
            QPushButton { background-color: #1877f2; color: white; font-weight: bold; font-size: 15px; border-radius: 6px; padding: 6px 12px; }
            QPushButton:hover { background-color: #166fe5; }
            QTableWidget { border: 2px solid #ccd0d5; border-radius: 10px; gridline-color: #dfe3ee; background-color: #f5f6f7; selection-background-color: #1877f2; selection-color: white; }
            QHeaderView::section { background-color: #e4e6eb; font-weight: bold; color: #050505; padding: 8px; border: none; }
            QTableWidget::item { padding: 10px; color: #000000; font-weight: bold; }
        """)

        self.docs = fetch_exporter_docs()
        self.edit_windows = []
        self.initUI()

    def initUI(self):
        layout = QVBoxLayout()
        layout.setAlignment(Qt.AlignmentFlag.AlignTop)

        # Header
        header_layout = QHBoxLayout()
        header_layout.setAlignment(Qt.AlignmentFlag.AlignCenter)

        logo = QLabel()
        pixmap = QPixmap("wwe.jpg")
        if not pixmap.isNull():
            pixmap = pixmap.scaled(200, 200, Qt.AspectRatioMode.KeepAspectRatio, Qt.TransformationMode.SmoothTransformation)
            logo.setPixmap(pixmap)
        else:
            logo.setText("No Logo")
        logo.setAlignment(Qt.AlignmentFlag.AlignCenter)
        header_layout.addWidget(logo)

        info_layout = QVBoxLayout()
        info_layout.setAlignment(Qt.AlignmentFlag.AlignCenter)
        name = QLabel("ZAKA Controls & Devices")
        name.setFont(QFont("Arial", 26, QFont.Weight.Bold))
        info_layout.addWidget(name)
        header_layout.addLayout(info_layout)
        layout.addLayout(header_layout)
        layout.addSpacing(20)

        # Table
        self.table = QTableWidget()
        self.table.setColumnCount(8)
        self.table.setHorizontalHeaderLabels([
            "exporter_ref", "iec", "order_no", "order_date", "tax_no",
            "Edit", "Hide", "Print"
        ])
        self.table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
        self.table.verticalHeader().setVisible(False)
        self.table.setAlternatingRowColors(True)
        self.table.setFont(QFont("Arial", 12, QFont.Weight.Bold))

        self.load_table_data()
        layout.addWidget(self.table)
        self.setLayout(layout)

    def load_table_data(self):
        self.table.setRowCount(0)
        for doc in self.docs:
            row_position = self.table.rowCount()
            self.table.insertRow(row_position)

            self.table.setItem(row_position, 0, QTableWidgetItem(str(doc.get("exporter_ref", ""))))
            self.table.setItem(row_position, 1, QTableWidgetItem(str(doc.get("iec", ""))))
            self.table.setItem(row_position, 2, QTableWidgetItem(str(doc.get("order_no", ""))))
            self.table.setItem(row_position, 3, QTableWidgetItem(str(doc.get("order_date", ""))))
            self.table.setItem(row_position, 4, QTableWidgetItem(str(doc.get("tax_no", ""))))

            # Buttons
            btn_edit = QPushButton("Edit")
            btn_edit.clicked.connect(lambda _, d=doc: self.open_edit_form(d))
            self.table.setCellWidget(row_position, 5, btn_edit)

            btn_hide = QPushButton("Hide")
            btn_hide.clicked.connect(lambda _, r=row_position: self.hide_row(r))
            self.table.setCellWidget(row_position, 6, btn_hide)

            btn_print = QPushButton("🖨 Print")
            btn_print.clicked.connect(lambda _, doc_id=str(doc["_id"]): self.print_doc(doc_id))
            self.table.setCellWidget(row_position, 7, btn_print)

            self.table.setRowHeight(row_position, 70)

    def open_edit_form(self, doc):
        edit_window = ExportFormEdit(doc)
        self.edit_windows.append(edit_window)
        edit_window.show()

    def hide_row(self, row):
        self.table.hideRow(row)
        QMessageBox.information(self, "Hidden", "Row has been hidden from the table view.")

    # ---------------- Print Function ----------------
    def print_doc(self, doc_id):
        try:
            client = MongoClient(MONGO_URL)
            db = client[DB_NAME]
            collection = db[COLLECTION_NAME]
            doc = collection.find_one({"_id": ObjectId(doc_id)})

            if not doc:
                QMessageBox.warning(self, "Not found", f"No document found with ID: {doc_id}")
                return

            # Map Mongo keys to DOCX placeholders
            mapped_data = {}
            for k, v in doc.items():
                if k == "_id":
                    continue
                new_key = PLACEHOLDER_MAPPING.get(k, k)
                mapped_data[new_key] = str(v)

            # Save JSON
            json_path = os.path.join(os.path.dirname(__file__), "data_one.json")
            with open(json_path, "w", encoding="utf-8") as f:
                json.dump(mapped_data, f, ensure_ascii=False, indent=4)

            # Generate DOCX


            # Get the full path of the template in the current directory
            template_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "1.docx")

            output_docx = os.path.join(os.path.dirname(__file__), "output_filled_1.docx")
            replace_placeholders(template_path, mapped_data, output_docx)

            # Convert DOCX -> Image -> PDF
            output_image = os.path.join(os.path.dirname(__file__), "output_filled_1.png")
            docx_to_image(output_docx, output_image)
            output_pdf = os.path.join(os.path.dirname(__file__), "output_filled_1.pdf")
            image_to_pdf(output_image, output_pdf)

            QMessageBox.information(
                self,
                "Success",
                f"✅ JSON, DOCX, and PDF generated successfully!\nJSON: {json_path}\nDOCX: {output_docx}\nPDF: {output_pdf}"
            )

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to generate documents: {str(e)}")


# ---------------- Run App ----------------
if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = ExportTable()
    window.show()
    sys.exit(app.exec())
