import re
import json
from docx import Document
from pydocx import PyDocX
from xhtml2pdf import pisa
import os
import collections
import collections.abc
from docx2pdf import convert as docx_to_pdf_direct
from PIL import Image
from docx2python import docx2python
if not hasattr(collections, 'Hashable'):
    collections.Hashable = collections.abc.Hashable
def extract_placeholders(docx_path):
    """Extract placeholders in the format «key» from a .docx file."""
    doc = Document(docx_path)
    placeholders = set()
    pattern = r'«(.*?)»'

    for para in doc.paragraphs:
        matches = re.findall(pattern, para.text)
        placeholders.update(matches)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    matches = re.findall(pattern, para.text)
                    placeholders.update(matches)

    return list(placeholders)


def replace_placeholders(docx_path, data, output_docx):
    """Replace placeholders with data values and save new .docx."""
    doc = Document(docx_path)
    pattern = r'«(.*?)»'

    def replace_in_para(para):
        for run in para.runs:
            matches = re.findall(pattern, run.text)
            for match in matches:
                if match in data:
                    run.text = run.text.replace(f'«{match}»', str(data[match]))

    for para in doc.paragraphs:
        replace_in_para(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_para(para)

    doc.save(output_docx)




def docx_to_image(docx_path, image_path):
    """Convert DOCX to image using python-docx + Pillow."""
    from docx import Document
    import tempfile

    doc = Document(docx_path)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"

    # Create an image with white background
    from PIL import Image, ImageDraw, ImageFont
    img = Image.new("RGB", (1200, 1600), color="white")
    draw = ImageDraw.Draw(img)

    # Use default PIL font
    font = ImageFont.load_default()
    draw.text((50, 50), text, fill="black", font=font)

    img.save(image_path)
    print(f"✅ Image saved: {image_path}")
    return image_path

def image_to_pdf(image_path, pdf_path):
    """Convert image to PDF using Pillow."""
    img = Image.open(image_path)
    img = img.convert("RGB")  # Ensure it's RGB
    img.save(pdf_path)
    print(f"✅ PDF from image saved: {pdf_path}")




template_path = "1.docx"
json_path = "data_one.json"

    # Load data from JSON
with open(json_path, "r", encoding="utf-8") as f:
    data = json.load(f)

    # Extract placeholders
keys = extract_placeholders(template_path)
print("🧩 Placeholders found in document:")
for k in keys:
    print(f"- {k}")

    # Output file paths
output_docx = "output_filled_1.docx"
output_pdf = "output_filled_1.pdf"

    # Replace placeholders
print("\n🔄 Replacing placeholders...")
# After replacing placeholders
replace_placeholders(template_path, data, output_docx)

# 1️⃣ Convert DOCX to image
output_image = "output_filled_1.png"
docx_to_image(output_docx, output_image)

# 2️⃣ Convert that image to PDF
image_pdf = "output_filled_1_from_image.pdf"
image_to_pdf(output_image, image_pdf)



