# sub_merge.py
import json
from docx import Document
import os
import re

def perform_mail_merge(json_path, template_path, output_path):
    """
    Simple mail merge: replaces placeholders in template docx with data from JSON.
    """
    if not os.path.exists(json_path):
        print(f"❌ {json_path} not found.")
        return
    if not os.path.exists(template_path):
        print(f"❌ {template_path} not found.")
        return

    # Load JSON data
    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    doc = Document(template_path)
    pattern = r'«(.*?)»'

    # Replace placeholders in paragraphs
    for para in doc.paragraphs:
        for run in para.runs:
            matches = re.findall(pattern, run.text)
            for key in matches:
                if key in data:
                    run.text = run.text.replace(f'«{key}»', str(data[key]))

    # Replace placeholders in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        matches = re.findall(pattern, run.text)
                        for key in matches:
                            if key in data:
                                run.text = run.text.replace(f'«{key}»', str(data[key]))

    doc.save(output_path)
    print(f"✅ Mail merge complete — saved to: {output_path}")
